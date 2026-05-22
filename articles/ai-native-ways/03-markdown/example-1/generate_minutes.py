#!/usr/bin/env python3
"""12 ヶ月分のサンプル議事録 .docx を生成する。

実際の議事録に近い構造(出席者・議題・決定事項・宿題)を持たせ、
書式(見出しスタイル、太字、箇条書き)も入れる。これをあとで pandoc で
Markdown に変換し、サイズ比とトークン比を測る。
"""
from __future__ import annotations

import random
from pathlib import Path

from docx import Document
from docx.shared import Pt

random.seed(42)

OUT = Path(__file__).parent / "docx"
OUT.mkdir(exist_ok=True)

ATTENDEES = ["山田", "鈴木", "高橋", "佐藤", "田中", "渡辺", "中村", "小林"]
TOPICS = [
    ("2026 年度予算案", "予算配分を部門ごとに見直す"),
    ("新製品ロードマップ", "Q2 にプロトタイプ完成を目指す"),
    ("採用計画", "エンジニア 3 名、デザイナー 1 名を募集"),
    ("オフィス移転", "賃料と通勤時間のバランスを再検討"),
    ("セキュリティ監査", "外部ベンダーによる監査を年内に実施"),
    ("AI 活用方針", "Claude を全社で導入、Office 依存を削減"),
    ("リモートワーク制度", "週 3 日のハイブリッドを正式化"),
    ("社内勉強会", "月 1 回の Markdown / Python 勉強会を開始"),
    ("品質保証プロセス", "テスト自動化率を 80% 以上に"),
    ("顧客フィードバック", "上半期の NPS を +12 ポイント改善"),
    ("コスト削減", "SaaS 利用料を 20% 削減目標"),
    ("広報戦略", "技術ブログを月 4 本ペースで継続"),
]
ACTIONS = [
    "次回までに資料を準備する",
    "関係部署にヒアリングする",
    "見積もりを 3 社から取得する",
    "詳細な工数見積もりを出す",
    "リスク一覧を整理する",
    "ステークホルダに事前共有する",
]
DECISIONS = [
    "本案で進めることを決定",
    "条件付きで承認、来月再確認",
    "再検討のため次回に持ち越し",
    "予算を 10% 増額して実施",
    "外部委託せず内製で進める",
]


def make_doc(month: int) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    doc.add_heading(f"2026 年 {month:02d} 月度 定例会議 議事録", level=0)
    doc.add_paragraph(f"日時: 2026 年 {month:02d} 月 15 日 14:00 - 16:00")
    doc.add_paragraph(f"場所: 本社 会議室 A")

    attendees = random.sample(ATTENDEES, k=random.randint(4, 6))
    p = doc.add_paragraph()
    p.add_run("出席者: ").bold = True
    p.add_run("、".join(attendees))

    doc.add_heading("議題", level=1)
    topics = random.sample(TOPICS, k=4)
    for i, (title, desc) in enumerate(topics, 1):
        doc.add_heading(f"{i}. {title}", level=2)
        doc.add_paragraph(desc)

        doc.add_heading("討議内容", level=3)
        for _ in range(random.randint(3, 5)):
            doc.add_paragraph(
                random.choice(
                    [
                        "現場からは慎重論と推進論が出た。",
                        "コスト面の懸念が指摘された。",
                        "他社事例を参考にすべきとの声があった。",
                        "段階的な導入が現実的との結論。",
                        "決定権者の参加を次回に求める。",
                        "前年同期比で成果を測定する方針。",
                    ]
                ),
                style="List Bullet",
            )

        doc.add_heading("決定事項", level=3)
        decision = random.choice(DECISIONS)
        p = doc.add_paragraph()
        p.add_run("決定: ").bold = True
        p.add_run(decision)

        doc.add_heading("宿題", level=3)
        for _ in range(random.randint(1, 2)):
            owner = random.choice(attendees)
            action = random.choice(ACTIONS)
            doc.add_paragraph(f"{owner}: {action}", style="List Number")

    doc.add_heading("次回開催", level=1)
    doc.add_paragraph(f"2026 年 {month + 1:02d} 月 15 日 14:00 -")

    return doc


def main():
    for month in range(1, 13):
        doc = make_doc(month)
        path = OUT / f"2026-{month:02d}-minutes.docx"
        doc.save(path)
        print(f"  {path.name}  ({path.stat().st_size:>6} bytes)")
    total = sum(p.stat().st_size for p in OUT.glob("*.docx"))
    print(f"\n合計: {total:,} bytes ({total / 1024:.1f} KB)")


if __name__ == "__main__":
    main()
